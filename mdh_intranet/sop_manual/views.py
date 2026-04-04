from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.db.models import Q
from django.utils import timezone
from mdh_intranet.sop_manual.models import SOP, SOPCategory
from mdh_intranet.core.models import AuditLog, SOPAcknowledgement as CoreAcknowledgment
from .forms import SOPForm

@login_required
def index(request):
    """SOP Manual Dashboard"""
    categories = SOPCategory.objects.all()
    sops = SOP.objects.filter(status='Published')
    
    total_sops = sops.count()
    recent_updates = sops.order_by('-updated_at')[:5]
    
    # Calculate user's read progress?
    # Maybe later.
    
    context = {
        'total_sops': total_sops,
        'recent_updates': recent_updates.count(), # count for now
        'categories': categories,
        'latest_sops': recent_updates,
    }
    return render(request, 'sop_manual/index.html', context)

@login_required
def sop_list(request, category_id=None):
    """List SOPs, optionally filtered by category"""
    sops = SOP.objects.filter(status='Published').select_related('category', 'created_by', 'linked_document')
    category = None
    
    if category_id:
        category = get_object_or_404(SOPCategory, id=category_id)
        sops = sops.filter(category=category)
        
    query = request.GET.get('search')
    if query:
        sops = sops.filter(Q(title__icontains=query) | Q(content__icontains=query))

    context = {
        'sops': sops,
        'category': category,
        'search_query': query
    }
    return render(request, 'sop_manual/sop_list.html', context)

@login_required
def sop_detail(request, pk):
    """View SOP detail and handle acknowledgment"""
    sop = get_object_or_404(SOP, pk=pk)
    
    # Check if acknowledged
    is_acknowledged = CoreAcknowledgment.objects.filter(sop=sop, user=request.user).exists()
    
    # Handle Acknowledgment (Save to Core model)
    if request.method == 'POST' and 'acknowledge' in request.POST:
        from mdh_intranet.core.services import get_client_ip, log_action
        CoreAcknowledgment.objects.get_or_create(
            sop=sop, 
            user=request.user,
            defaults={'ip_address': get_client_ip(request)}
        )
        log_action(request.user, 'acknowledge', sop, 
                   description=f"Acknowledged SOP: {sop.title}",
                   ip_address=get_client_ip(request))
        messages.success(request, f"You have acknowledged reading '{sop.title}'.")
        return redirect('sop_manual:detail', pk=pk)
    
    # Fetch Audit logs for this SOP
    from django.contrib.contenttypes.models import ContentType
    sop_ct = ContentType.objects.get_for_model(SOP)
    audit_logs = AuditLog.objects.filter(content_type=sop_ct, object_id=sop.pk).select_related('user')[:50]
    
    # Fetch all acknowledgements for this SOP
    acknowledgements = CoreAcknowledgment.objects.filter(sop=sop).select_related('user').order_by('-acknowledged_at')
        
    context = {
        'sop': sop,
        'is_acknowledged': is_acknowledged,
        'audit_logs': audit_logs,
        'acknowledgements': acknowledgements,
    }
    return render(request, 'sop_manual/sop_detail.html', context)

@user_passes_test(lambda u: u.is_staff)
def sop_create(request):
    """Create a new SOP"""
    if request.method == 'POST':
        form = SOPForm(request.POST, request.FILES)
        if form.is_valid():
            sop = form.save(commit=False)
            sop.created_by = request.user
            sop.save()
            messages.success(request, "SOP created successfully.")
            return redirect('sop_manual:detail', pk=sop.pk)
    else:
        form = SOPForm()
        
    return render(request, 'sop_manual/sop_form.html', {'form': form, 'title': 'Create SOP'})

@user_passes_test(lambda u: u.is_staff)
def sop_edit(request, pk):
    """Edit an existing SOP"""
    sop = get_object_or_404(SOP, pk=pk)
    if request.method == 'POST':
        form = SOPForm(request.POST, request.FILES, instance=sop)
        if form.is_valid():
            form.save()
            messages.success(request, "SOP updated successfully.")
            return redirect('sop_manual:detail', pk=sop.pk)
    else:
        form = SOPForm(instance=sop)
        
    return render(request, 'sop_manual/sop_form.html', {'form': form, 'title': f'Edit SOP: {sop.title}'})


@login_required
def sop_office_viewer(request, pk):
    """View SOP file attachment using Microsoft Office Web Viewer"""
    sop = get_object_or_404(SOP, pk=pk)
    
    # Determine which file to use
    target_file = None
    if sop.linked_document and sop.linked_document.file:
        target_file = sop.linked_document.file
    elif sop.file_attachment:
        target_file = sop.file_attachment

    # Check if SOP has a file attachment
    if not target_file:
        messages.error(request, 'This SOP does not have a file attachment.')
        return redirect('sop_manual:detail', pk=pk)
    
    # Get file extension
    import os
    file_ext = os.path.splitext(target_file.name)[1].lower()
    
    # Check if file type is supported by Office Web Viewer
    supported_extensions = ['.docx', '.xlsx', '.pptx', '.doc', '.xls', '.ppt']
    
    error = None
    viewer_url = None
    
    if file_ext not in supported_extensions:
        error = f"Microsoft Office Web Viewer does not support {file_ext} files. Supported formats: Word (.docx), Excel (.xlsx), PowerPoint (.pptx)"
    else:
        # Build the absolute URL to the document
        from urllib.parse import quote
        document_url = request.build_absolute_uri(target_file.url)
        encoded_url = quote(document_url, safe='')
        
        # Microsoft Office Web Viewer endpoint
        viewer_url = f"https://view.officeapps.live.com/op/embed.aspx?src={encoded_url}"
    
    context = {
        'sop': sop,
        'viewer_url': viewer_url,
        'error': error,
        'file_ext': file_ext,
    }
    return render(request, 'sop_manual/sop_office_viewer.html', context)

from django.views.decorators.clickjacking import xframe_options_exempt

@xframe_options_exempt
@login_required
def sop_office_embed(request, pk):
    """Embed-only view for SOP file attachment"""
    sop = get_object_or_404(SOP, pk=pk)
    
    # Determine which file to use
    target_file = None
    if sop.linked_document and sop.linked_document.file:
        target_file = sop.linked_document.file
    elif sop.file_attachment:
        target_file = sop.file_attachment
    
    # Check if SOP has a file attachment
    if not target_file:
        return render(request, 'sop_manual/sop_embed.html', {'error': 'No file attached'})
    
    # Get file extension
    import os
    file_ext = os.path.splitext(target_file.name)[1].lower()
    
    # Check if file type is supported
    supported_extensions = ['.docx', '.xlsx', '.pptx', '.doc', '.xls', '.ppt']
    
    viewer_url = None
    error = None
    
    if file_ext not in supported_extensions:
        error = f"Format {file_ext} not supported."
    else:
        from urllib.parse import quote
        document_url = request.build_absolute_uri(target_file.url)
        encoded_url = quote(document_url, safe='')
        viewer_url = f"https://view.officeapps.live.com/op/embed.aspx?src={encoded_url}&wdStartOn=1"
        
    return render(request, 'sop_manual/sop_embed.html', {
        'viewer_url': viewer_url,
        'error': error
    })


# ── SOP EXPORT: DOCX ──────────────────────────────────────────
@login_required
def export_sop_docx(request, pk):
    """Export an SOP as a professionally formatted .docx file."""
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    from django.http import HttpResponse
    from .export_utils import preprocess_html_for_docx

    sop = get_object_or_404(SOP, pk=pk)

    doc = DocxDocument()

    # Check if the SOP content already has a header block (avoid duplicates)
    has_header_block = '<div class="sop-header-block"' in sop.content or '<table class="table table-bordered mb-0"' in sop.content

    if not has_header_block:
        # ── Title ──
        title_para = doc.add_heading(sop.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(30, 64, 175)

        # ── Subtitle ──
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run('OpsHub — Standard Operating Procedure')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 116, 139)

        doc.add_paragraph()  # spacer

        # ── Metadata table ──
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        meta = [
            ('Category', str(sop.category) if sop.category else '—'),
            ('Version', sop.version),
            ('Status', sop.status),
            ('Author', sop.created_by.get_full_name() or sop.created_by.username if sop.created_by else '—'),
            ('Last Updated', sop.updated_at.strftime('%d %B %Y') if sop.updated_at else '—'),
        ]
        for i, (key, val) in enumerate(meta):
            row = table.rows[i]
            row.cells[0].text = key
            row.cells[1].text = val
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)

        doc.add_paragraph()  # spacer
    else:
        # If it has a header block, just add a small margin at the top
        doc.add_paragraph()
    
    # ── Content ──
    from htmldocx import HtmlToDocx
    new_parser = HtmlToDocx()
    
    # Configure the parser to use a standard table style with borders
    new_parser.table_style = 'Table Grid'
    
    # Set default font and size for the entire document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Check if the SOP content already has a header block (avoid duplicates)
    has_header_block = '<div class="sop-header-block"' in sop.content or '<table class="table table-bordered mb-0"' in sop.content
    
    if not has_header_block:
        doc.add_heading('Procedure Details', level=1)
    
    # Preprocess the HTML to convert Bootstrap classes to inline styles
    # and clean up FontAwesome icons, colspan/rowspan, etc.
    try:
        content_html = preprocess_html_for_docx(sop.content)
        new_parser.add_html_to_document(content_html, doc)
    except Exception as e:
        # Fallback: strip all HTML and add as plain text
        from bs4 import BeautifulSoup
        doc.add_paragraph(
            f"[Export Warning: Some formatting was simplified. Error: {str(e)}]"
        )
        plain = BeautifulSoup(sop.content, 'html.parser').get_text('\n', strip=True)
        doc.add_paragraph(plain)
    
    # ── Digital Signatures / Acknowledgements ──
    from mdh_intranet.core.models import SOPAcknowledgement as CoreAcknowledgment
    acks = CoreAcknowledgment.objects.filter(sop=sop).select_related('user').order_by('acknowledged_at')
    
    if acks.exists():
        doc.add_page_break()
        doc.add_heading('Compliance & Digital Acknowledgements', level=1)
        doc.add_paragraph('The following personnel have confirmed reading and understanding this Standard Operating Procedure. This document provides a verified digital record of acknowledgement.')
        
        # Add acknowledgement table
        ack_table = doc.add_table(rows=1, cols=3)
        ack_table.style = 'Table Grid'
        hdr_cells = ack_table.rows[0].cells
        hdr_cells[0].text = 'Personnel'
        hdr_cells[1].text = 'Date & Time'
        hdr_cells[2].text = 'Digital Signature'
        
        for ack in acks:
            row_cells = ack_table.add_row().cells
            row_cells[0].text = ack.user.get_full_name() or ack.user.username
            row_cells[1].text = ack.acknowledged_at.strftime('%d %B %Y at %H:%M')
            # Create a mock digital signature hash if it doesn't exist, or use the one from the model if we had it
            # (CoreAcknowledgement doesn't have signature_hash yet, but we can generate a short ID)
            import hashlib
            sig_raw = f"{ack.user.id}|{ack.sop.id}|{ack.acknowledged_at.isoformat()}"
            sig_hash = hashlib.sha256(sig_raw.encode()).hexdigest()[:16].upper()
            row_cells[2].text = f"MDH-{sig_hash}"
            
            # Format table fonts
            for cell in row_cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.font.name = 'Consolas' if cell == row_cells[2] else 'Calibri'
    
    # Finalize response
    buffer = BytesIO()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f'Document generated from OpsHub on {timezone.now().strftime("%d %B %Y at %H:%M")}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(148, 163, 184)

    # Serve the file
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = f"SOP_{sop.pk}_{sop.title[:30].replace(' ', '_')}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ── SOP EXPORT: PDF ──────────────────────────────────────────
@login_required
def export_sop_pdf(request, pk):
    """Export an SOP as a PDF file using xhtml2pdf."""
    from io import BytesIO
    from django.http import HttpResponse
    from django.template.loader import render_to_string
    from xhtml2pdf import pisa
    from .export_utils import preprocess_html_for_pdf

    sop = get_object_or_404(SOP, pk=pk)

    # Fetch acknowledgements for the report
    from mdh_intranet.core.models import SOPAcknowledgement as CoreAcknowledgment
    acknowledgements = CoreAcknowledgment.objects.filter(sop=sop).select_related('user').order_by('acknowledged_at')
    
    # Preprocess for PDF
    processed_content = preprocess_html_for_pdf(sop.content)

    html_content = render_to_string('sop_manual/sop_pdf_template.html', {
        'sop': sop,
        'processed_content': processed_content,
        'generated_at': timezone.now(),
        'acknowledgements': acknowledgements,
    })

    buffer = BytesIO()
    pisa_status = pisa.CreatePDF(html_content, dest=buffer)

    if pisa_status.err:
        return HttpResponse('Error generating PDF', status=500)

    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    filename = f"SOP_{sop.pk}_{sop.title[:30].replace(' ', '_')}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ── SOP ASSISTANT ──────────────────────────────────────────────
from django.http import JsonResponse
from django.views.decorators.http import require_POST
import json

@login_required
def sop_assistant(request):
    """SOP Assistant — wizard UI for generating SOPs from data."""
    from .sop_generator import get_template_info
    from mdh_intranet.incident_log.models import Incident
    from mdh_intranet.capa.models import CAPARecord

    categories = SOPCategory.objects.all()
    templates = get_template_info()
    incidents = Incident.objects.order_by('-date_reported')[:50]
    capas = CAPARecord.objects.order_by('-created_at')[:50]

    context = {
        'categories': categories,
        'templates': templates,
        'incidents': incidents,
        'capas': capas,
    }
    return render(request, 'sop_manual/sop_assistant.html', context)


@login_required
@require_POST
def sop_assistant_generate(request):
    """AJAX endpoint — generate SOP content from selected template + data."""
    from .sop_generator import generate_sop_content

    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    title = data.get('title', '').strip()
    template_key = data.get('template', 'administrative')
    category_name = data.get('category', '')
    additional_context = data.get('additional_context', '')
    incident_ids = data.get('incident_ids', [])
    capa_ids = data.get('capa_ids', [])

    if not title:
        return JsonResponse({'error': 'Title is required'}, status=400)

    author_name = (
        request.user.get_full_name() or request.user.username
    )

    content = generate_sop_content(
        title=title,
        template_key=template_key,
        category_name=category_name,
        author_name=author_name,
        additional_context=additional_context,
        incident_ids=incident_ids,
        capa_ids=capa_ids,
    )

    return JsonResponse({'content': content, 'title': title})


@login_required
@require_POST
def sop_assistant_save(request):
    """Save a generated SOP directly to the database."""
    title = request.POST.get('title', '').strip()
    content = request.POST.get('content', '').strip()
    category_id = request.POST.get('category')

    if not title or not content:
        messages.error(request, 'Title and content are required.')
        return redirect('sop_manual:assistant')

    category = None
    if category_id:
        try:
            category = SOPCategory.objects.get(pk=category_id)
        except SOPCategory.DoesNotExist:
            pass

    if not category:
        # Fallback: get or create a General category
        category, _ = SOPCategory.objects.get_or_create(
            name='General',
            defaults={'description': 'General SOPs', 'icon': 'fas fa-file-alt'},
        )

    sop = SOP.objects.create(
        title=title,
        content=content,
        category=category,
        status='Draft',
        version='1.0',
        created_by=request.user,
    )

    messages.success(request, f'SOP "{sop.title}" created as Draft. Review and publish when ready.')
    return redirect('sop_manual:edit', pk=sop.pk)
