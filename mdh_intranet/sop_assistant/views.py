import json
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.http import JsonResponse
from django.db.models import Q
from django.utils import timezone

from .models import SOPTemplate, SOPDraft, SOPDraftSection, ValidationResult
from .forms import DraftMetadataForm, SectionContentForm, ICDCodeForm
from .validators import SOPValidator, compile_sop_html
from mdh_intranet.icd11_tools.models import ICDCode
from mdh_intranet.sop_manual.models import SOP, SOPCategory


@login_required
def index(request):
    """SOP Assistant Dashboard — overview of drafts and templates."""
    templates = SOPTemplate.objects.filter(is_active=True)
    my_drafts = SOPDraft.objects.filter(author=request.user).exclude(status='discarded')
    
    # Stats
    total_drafts = my_drafts.count()
    in_progress = my_drafts.filter(status__in=['template_selected', 'drafting', 'icd_review']).count()
    validated = my_drafts.filter(status__in=['validated', 'ready']).count()
    published = my_drafts.filter(status='published').count()

    context = {
        'templates': templates,
        'my_drafts': my_drafts[:10],
        'total_drafts': total_drafts,
        'in_progress': in_progress,
        'validated': validated,
        'published': published,
    }
    return render(request, 'sop_assistant/index.html', context)


@login_required
def select_template(request):
    """Step 1: Choose a template to start drafting."""
    templates = SOPTemplate.objects.filter(is_active=True)
    
    # Filter by category if provided
    category_filter = request.GET.get('category', '')
    if category_filter:
        templates = templates.filter(category=category_filter)

    categories = SOPTemplate.CATEGORY_CHOICES

    if request.method == 'POST':
        template_id = request.POST.get('template_id')
        if template_id:
            template = get_object_or_404(SOPTemplate, pk=template_id, is_active=True)
            
            # Create a new draft
            draft = SOPDraft.objects.create(
                title='',
                template=template,
                author=request.user,
                status='template_selected',
            )
            
            # Create section entries from template
            for i, section in enumerate(template.sections):
                SOPDraftSection.objects.create(
                    draft=draft,
                    section_key=section['key'],
                    section_label=section['label'],
                    order=i,
                    is_required=section.get('required', True),
                )

            # Increment usage
            template.usage_count += 1
            template.save()

            messages.success(request, f'Draft started with template: {template.name}')
            return redirect('sop_assistant:draft_metadata', draft_id=draft.pk)

    context = {
        'templates': templates,
        'categories': categories,
        'selected_category': category_filter,
    }
    return render(request, 'sop_assistant/select_template.html', context)


@login_required
def draft_metadata(request, draft_id):
    """Step 2: Fill in title, category, version."""
    draft = get_object_or_404(SOPDraft, pk=draft_id, author=request.user)

    from mdh_intranet.incident_log.models import Incident
    from mdh_intranet.capa.models import CAPARecord

    if request.method == 'POST':
        form = DraftMetadataForm(request.POST, instance=draft)
        if form.is_valid():
            # Standard fields + M2M fields
            draft = form.save(commit=False)
            draft.status = 'drafting'
            draft.save()
            form.save_m2m() # Required for many-to-many fields when commit=False
            
            messages.success(request, 'Metadata and Intelligence Context saved. Now fill in the SOP content.')
            return redirect('sop_assistant:draft_content', draft_id=draft.pk)
    else:
        form = DraftMetadataForm(instance=draft)
        # Limit options to recent relevant items
        form.fields['referenced_incidents'].queryset = Incident.objects.all().order_by('-created_at')[:20]
        form.fields['referenced_capas'].queryset = CAPARecord.objects.all().order_by('-created_at')[:20]

    context = {
        'draft': draft,
        'form': form,
        'step': 2,
        'total_steps': 5 if draft.is_clinical else 4,
    }
    return render(request, 'sop_assistant/draft_metadata.html', context)


@login_required
def draft_content(request, draft_id):
    """Step 3: Fill in template sections."""
    draft = get_object_or_404(SOPDraft, pk=draft_id, author=request.user)
    sections = draft.sections.all().order_by('order')

    if request.method == 'POST':
        # Save each section
        for section in sections:
            field_key = f"section_{section.section_key}"
            content = request.POST.get(field_key, '')
            section.content = content
            section.save()

        draft.status = 'drafting'
        draft.save()

        # Determine next step
        if draft.is_clinical:
            messages.success(request, 'Content saved. Add ICD-11 codes for this clinical SOP.')
            return redirect('sop_assistant:draft_icd', draft_id=draft.pk)
        else:
            messages.success(request, 'Content saved. Ready for validation.')
            return redirect('sop_assistant:draft_validate', draft_id=draft.pk)

    # Pre-fill form with existing content
    initial = {}
    for section in sections:
        initial[f"section_{section.section_key}"] = section.content

    context = {
        'draft': draft,
        'sections': sections,
        'initial_data': initial,
        'step': 3,
        'total_steps': 5 if draft.is_clinical else 4,
    }
    return render(request, 'sop_assistant/draft_content.html', context)


@login_required
def draft_icd(request, draft_id):
    """Step 4 (clinical only): ICD-11 code lookup and attachment."""
    draft = get_object_or_404(SOPDraft, pk=draft_id, author=request.user)

    if not draft.is_clinical:
        return redirect('sop_assistant:draft_validate', draft_id=draft.pk)

    if request.method == 'POST':
        action = request.POST.get('action', '')

        if action == 'add_code':
            code_id = request.POST.get('code_id')
            if code_id:
                try:
                    icd = ICDCode.objects.get(pk=code_id)
                    codes = draft.icd_codes
                    # Avoid duplicates
                    if not any(c['code'] == icd.code for c in codes):
                        codes.append({'code': icd.code, 'description': icd.description})
                        draft.icd_codes = codes
                        draft.save()
                        messages.success(request, f'Added ICD-11 code: {icd.code}')
                except ICDCode.DoesNotExist:
                    messages.warning(request, 'Code not found.')

        elif action == 'remove_code':
            code_str = request.POST.get('code_str')
            codes = draft.icd_codes
            codes = [c for c in codes if c['code'] != code_str]
            draft.icd_codes = codes
            draft.save()
            messages.info(request, f'Removed code: {code_str}')

        elif action == 'continue':
            draft.status = 'icd_review'
            draft.save()
            return redirect('sop_assistant:draft_validate', draft_id=draft.pk)

    # Handle search
    search_query = request.GET.get('search', '')
    search_results = None
    if search_query:
        search_results = ICDCode.objects.filter(
            Q(code__icontains=search_query) |
            Q(description__icontains=search_query)
        )[:20]

    context = {
        'draft': draft,
        'search_query': search_query,
        'search_results': search_results,
        'attached_codes': draft.icd_codes,
        'step': 4,
        'total_steps': 5,
    }
    return render(request, 'sop_assistant/draft_icd.html', context)


@login_required
def draft_validate(request, draft_id):
    """Step 4/5: Validation and review."""
    draft = get_object_or_404(SOPDraft, pk=draft_id, author=request.user)

    if request.method == 'POST':
        action = request.POST.get('action', '')

        if action == 'validate':
            validator = SOPValidator(draft)
            score = validator.validate_all()
            messages.info(request, f'Validation complete. Score: {score}/100')

        elif action == 'compile':
            # Compile the HTML
            html = compile_sop_html(draft)
            draft.compiled_content = html
            draft.status = 'ready'
            draft.save()
            messages.success(request, 'SOP compiled and ready for publication!')
            return redirect('sop_assistant:draft_preview', draft_id=draft.pk)

    validations = draft.validations.all()
    errors = validations.filter(severity='error').count()
    warnings = validations.filter(severity='warning').count()
    passed = validations.filter(severity='success').count()

    context = {
        'draft': draft,
        'validations': validations,
        'errors': errors,
        'warnings': warnings,
        'passed': passed,
        'step': 5 if draft.is_clinical else 4,
        'total_steps': 5 if draft.is_clinical else 4,
    }
    return render(request, 'sop_assistant/draft_validate.html', context)


@login_required
def draft_preview(request, draft_id):
    """Preview the compiled SOP and optionally publish."""
    draft = get_object_or_404(SOPDraft, pk=draft_id, author=request.user)

    if not draft.compiled_content:
        html = compile_sop_html(draft)
        draft.compiled_content = html
        draft.save()

    if request.method == 'POST' and request.POST.get('action') == 'publish':
        # Create the SOP in the SOP Manual
        sop = SOP.objects.create(
            title=draft.title,
            category=draft.target_category,
            content=draft.compiled_content,
            version=draft.version,
            status='Published',
            created_by=draft.author,
        )
        draft.published_sop = sop
        draft.status = 'published'
        draft.save()

        messages.success(request, f'🎉 SOP "{sop.title}" published successfully to the SOP Manual!')
        return redirect('sop_manual:detail', pk=sop.pk)

    context = {
        'draft': draft,
    }
    return render(request, 'sop_assistant/draft_preview.html', context)


@login_required
def draft_discard(request, draft_id):
    """Discard a draft."""
    draft = get_object_or_404(SOPDraft, pk=draft_id, author=request.user)
    if request.method == 'POST':
        draft.status = 'discarded'
        draft.save()
        messages.info(request, 'Draft discarded.')
    return redirect('sop_assistant:index')


@login_required
def export_draft_docx(request, draft_id):
    """Export an SOP Draft as a professionally formatted .docx file."""
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    from django.http import HttpResponse
    from mdh_intranet.sop_manual.export_utils import preprocess_html_for_docx

    draft = get_object_or_404(SOPDraft, pk=draft_id, author=request.user)
    if not draft.compiled_content:
        draft.compiled_content = compile_sop_html(draft)
        draft.save()

    doc = DocxDocument()

    # ── Content ──
    from htmldocx import HtmlToDocx
    new_parser = HtmlToDocx()
    new_parser.table_style = 'Table Grid'
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    try:
        content_html = preprocess_html_for_docx(draft.compiled_content)
        new_parser.add_html_to_document(content_html, doc)
    except Exception as e:
        from bs4 import BeautifulSoup
        doc.add_paragraph(f"[Export Warning: Some formatting was simplified. Error: {str(e)}]")
        plain = BeautifulSoup(draft.compiled_content, 'html.parser').get_text('\n', strip=True)
        doc.add_paragraph(plain)

    # Finalize response
    buffer = BytesIO()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f'Draft Document generated from OpsHub on {timezone.now().strftime("%d %B %Y at %H:%M")}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(148, 163, 184)

    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = f"Draft_SOP_{draft.pk}_{draft.title[:30].replace(' ', '_')}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def export_draft_pdf(request, draft_id):
    """Export an SOP Draft as a PDF file using xhtml2pdf."""
    from io import BytesIO
    from django.http import HttpResponse
    from django.template.loader import render_to_string
    from xhtml2pdf import pisa
    from mdh_intranet.sop_manual.export_utils import preprocess_html_for_pdf

    draft = get_object_or_404(SOPDraft, pk=draft_id, author=request.user)
    if not draft.compiled_content:
        draft.compiled_content = compile_sop_html(draft)
        draft.save()

    # Create a mock SOP object for the template compatibility
    from types import SimpleNamespace
    mock_sop = SimpleNamespace(
        title=draft.title,
        version=draft.version,
        category=draft.target_category,
        created_by=draft.author,
        status="Draft",
        updated_at=draft.updated_at
    )

    processed_content = preprocess_html_for_pdf(draft.compiled_content)

    html_content = render_to_string('sop_manual/sop_pdf_template.html', {
        'sop': mock_sop,
        'processed_content': processed_content,
        'generated_at': timezone.now(),
        'acknowledgements': [],
    })

    buffer = BytesIO()
    pisa_status = pisa.CreatePDF(html_content, dest=buffer)

    if pisa_status.err:
        return HttpResponse('Error generating PDF', status=500)

    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    filename = f"Draft_SOP_{draft.pk}_{draft.title[:30].replace(' ', '_')}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def draft_detail(request, draft_id):
    """View draft status and details."""
    draft = get_object_or_404(SOPDraft, pk=draft_id, author=request.user)
    sections = draft.sections.all()
    validations = draft.validations.all()

    context = {
        'draft': draft,
        'sections': sections,
        'validations': validations,
    }
    return render(request, 'sop_assistant/draft_detail.html', context)


# ─── API Endpoints ───────────────────────────────────────────────

@login_required
def api_icd_search(request):
    """AJAX endpoint for ICD-11 code search."""
    query = request.GET.get('q', '').strip()
    if len(query) < 2:
        return JsonResponse({'results': []})

    codes = ICDCode.objects.filter(
        Q(code__icontains=query) | Q(description__icontains=query)
    )[:15]

    results = [
        {
            'id': code.pk,
            'code': code.code,
            'description': code.description,
            'chapter': code.chapter,
        }
        for code in codes
    ]
    return JsonResponse({'results': results})


@login_required
def api_autosave(request, draft_id):
    """AJAX endpoint for auto-saving draft sections."""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)

    draft = get_object_or_404(SOPDraft, pk=draft_id, author=request.user)

    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    section_key = data.get('section_key')
    content = data.get('content', '')

    try:
        section = SOPDraftSection.objects.get(draft=draft, section_key=section_key)
        section.content = content
        section.save()
        return JsonResponse({'status': 'saved', 'timestamp': timezone.now().isoformat()})
    except SOPDraftSection.DoesNotExist:
        return JsonResponse({'error': 'Section not found'}, status=404)


@login_required
def api_ai_suggest(request, draft_id):
    """
    Intelligent endpoint to suggest content for a specific section.
    Uses title, category, and hospital data (incidents/CAPAs) to generate a draft.
    """
    draft = get_object_or_404(SOPDraft, pk=draft_id)
    section_key = request.GET.get('section_key', '')
    
    if not section_key:
        return JsonResponse({'error': 'section_key required'}, status=400)

    # Use the shared generator logic if possible, or a local intelligence engine
    from mdh_intranet.sop_manual.sop_generator import SECTION_BUILDERS
    
    label = "the procedure"
    try:
        section = draft.sections.get(section_key=section_key)
        label = section.section_label
    except SOPDraftSection.DoesNotExist:
        pass

    content = ""
    if section_key in SECTION_BUILDERS:
        _, template_str = SECTION_BUILDERS[section_key]
        try:
            content = template_str.format(
                title=draft.title,
                title_lower=draft.title.lower(),
                category=draft.target_category.name if draft.target_category else "General",
                author=request.user.get_full_name() or request.user.username,
                date_str=timezone.now().strftime('%d %B %Y')
            )
        except Exception:
            content = template_str

    # High Intelligence: Inject incident/CAPA insights into specific sections if they exist
    findings = []
    for inc in draft.referenced_incidents.all()[:5]:
        if inc.corrective_actions:
            findings.append(f"<li><strong>INC-{inc.id}:</strong> {inc.corrective_actions}</li>")
        elif inc.immediate_action:
            findings.append(f"<li><strong>INC-{inc.id}:</strong> {inc.immediate_action}</li>")
            
    for capa in draft.referenced_capas.all()[:5]:
        if capa.preventive_action_plan:
            findings.append(f"<li><strong>CAPA-{capa.id}:</strong> {capa.preventive_action_plan}</li>")

    if findings:
        insights_html = (
            "<div class='alert alert-info mt-4'>"
            "<strong><i class='fas fa-lightbulb me-2'></i>AI Suggested Inclusions (from related Incidents/CAPAs):</strong>"
            "<ul>" + "".join(findings) + "</ul>"
            "</div>"
        )
        
        # Append insights intelligently based on the section
        if section_key in ('procedure_steps', 'safety_considerations', 'quality_control', 'prerequisites', 'training_requirements'):
            content += insights_html
        elif not content:
            content = f"<p>Regarding <strong>{label}</strong> for <em>{draft.title}</em>, ensure the following measures are taken based on recent events:</p>" + insights_html

    if not content:
        # Fallback for undefined sections
        content = (
            f"<p>Describe the <strong>{label}</strong> for <em>{draft.title}</em> here.</p>"
            "<ul><li>Ensure safety first</li><li>Verify patient identifiers</li><li>Document outcome</li></ul>"
        )

    return JsonResponse({'suggestion': content})


@login_required
def api_get_context(request, draft_id):
    """
    Fetch relevant context (incidents, CAPAs, SOPs) for the current draft title.
    """
    draft = get_object_or_404(SOPDraft, pk=draft_id)
    from mdh_intranet.sop_manual.sop_generator import gather_context_data
    
    context = gather_context_data(
        topic=draft.title,
        category_name=draft.target_category.name if draft.target_category else ""
    )
    
    # Add manually selected items to context if they aren't already there
    referenced_incidents = list(draft.referenced_incidents.values('pk', 'title', 'status', 'category'))
    referenced_capas = list(draft.referenced_capas.values('pk', 'title', 'status', 'priority'))
    
    context['manually_selected'] = {
        'incidents': referenced_incidents,
        'capas': referenced_capas
    }
    
    return JsonResponse(context)


@login_required
def api_validate(request, draft_id):
    """
    Run the validation engine and return results as JSON for real-time feedback.
    """
    draft = get_object_or_404(SOPDraft, pk=draft_id)
    from .validators import SOPValidator
    
    validator = SOPValidator(draft)
    score = validator.validate_all()
    
    validations = draft.validations.all().values(
        'rule_code', 'severity', 'message', 'suggestion'
    )
    
    return JsonResponse({
        'score': score,
        'results': list(validations)
    })


@login_required
def api_icd_suggest(request, draft_id):
    """
    Intelligently suggest ICD-11 codes based on the SOP title and category.
    """
    draft = get_object_or_404(SOPDraft, pk=draft_id)
    query = draft.title.lower()
    
    # Generic smart suggestions based on keywords, removing common SOP words
    stop_words = {'protocol', 'procedure', 'management', 'guideline', 'sop', 'policy', 'standard', 'care', 'routine'}
    keywords = [w for w in query.split() if len(w) > 3 and w not in stop_words]
    q = Q()
    for word in keywords[:4]:
        q |= Q(description__icontains=word) | Q(code__icontains=word)
    
    if q:
        suggestions = ICDCode.objects.filter(q)[:5]
    else:
        suggestions = ICDCode.objects.none()
    
    results = [
        {'id': s.pk, 'code': s.code, 'description': s.description}
        for s in suggestions
    ]
    
    return JsonResponse({'suggestions': results})
