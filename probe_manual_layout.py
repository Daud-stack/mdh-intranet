
import os
import django
from docx import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
import re

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'mdh_intranet.settings')
django.setup()
from mdh_intranet.documents.models import Document

def probe_sop_layout(doc_id, start_line):
    d = Document.objects.get(id=doc_id)
    doc = DocxDocument(d.file.path)
    
    # We'll look at a different SOP to see the variety
    # Let's check around paragraph 5000
    print(f"--- PROBING AROUND PARAGRAPH 5000 ---")
    for i in range(5000, 5200):
        if i < len(doc.paragraphs):
            t = doc.paragraphs[i].text.strip()
            if t:
                print(f"[{i}] {t}")
                if re.search(r'MDH-[A-Z]+-\d+', t):
                    print("!!! SOP HEADER FOUND !!!")

if __name__ == "__main__":
    probe_sop_layout(4, 5000)
