
import os
import django
from docx import Document as DocxDocument
import re
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'mdh_intranet.settings')
django.setup()
from mdh_intranet.documents.models import Document

def extract_full_sop(doc_id, start_idx):
    d = Document.objects.get(id=doc_id)
    doc = DocxDocument(d.file.path)
    content = []
    
    elements = list(doc.element.body.iterchildren())
    found_next = False
    
    for i in range(start_idx, len(elements)):
        element = elements[i]
        if isinstance(element, CT_P):
            p = Paragraph(element, doc)
            text = p.text.strip()
            # If we see another MDH-XXX-### after the first few paragraphs, it might be the next SOP
            if i > start_idx + 10 and re.search(r'MDH-[A-Z]+-\d+', text) and len(text) < 100:
                print(f"DEBUG: Found next SOP at {i}: {text}")
                break
            if text:
                content.append({"type": "para", "text": text})
        elif isinstance(element, CT_Tbl):
            t = Table(element, doc)
            data = [[c.text.strip() for c in r.cells] for r in t.rows]
            content.append({"type": "table", "data": data})
            
    return content

if __name__ == "__main__":
    # Start at element index (Note: body.iterchildren might not match paragraph index exactly)
    # The previous script used paragraph index. Paragraph index 182 might not be element index 182.
    
    d = Document.objects.get(id=4)
    doc = DocxDocument(d.file.path)
    target_p = doc.paragraphs[182]
    
    # Find element index of target_p
    elements = list(doc.element.body.iterchildren())
    start_elm_idx = -1
    for i, el in enumerate(elements):
        if isinstance(el, CT_P):
            if Paragraph(el, doc).text == target_p.text:
                # Double check context
                # This is a bit loose but okay for now
                start_elm_idx = i
                break
    
    if start_elm_idx != -1:
        sop_content = extract_full_sop(4, start_elm_idx)
        for item in sop_content:
            if item['type'] == 'para':
                print(f"P: {item['text']}")
            else:
                print(f"T: {item['data']}")
