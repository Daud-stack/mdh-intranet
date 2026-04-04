
import os
import django
from docx import Document as DocxDocument
import re

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'mdh_intranet.settings')
django.setup()
from mdh_intranet.documents.models import Document

def find_sop_start(doc_id, code):
    d = Document.objects.get(id=doc_id)
    doc = DocxDocument(d.file.path)
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if code in t and i > 100: # Skip TOC
            print(f"Potential Start at [{i}]: {t}")
            # Check next few paragraphs
            for j in range(i+1, i+10):
                if j < len(doc.paragraphs):
                    nt = doc.paragraphs[j].text.strip()
                    if nt:
                        print(f"  [{j}] {nt}")
            break

if __name__ == "__main__":
    find_sop_start(4, "MDH-FIN-001")
