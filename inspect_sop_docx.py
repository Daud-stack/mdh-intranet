
from docx import Document
import os

file_path = r"c:\Users\HomePC\mdh_intranet\media\documents\2026\02\MDH_SOP_Manual_Full_Content.docx"

try:
    doc = Document(file_path)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print("First 20 paragraphs:")
    for i, p in enumerate(doc.paragraphs[:20]):
        print(f"{i}: [{p.style.name}] {p.text[:100]}...")
except Exception as e:
    print(f"Error: {e}")
