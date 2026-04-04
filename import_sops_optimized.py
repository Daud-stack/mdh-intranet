
from docx import Document
import os
import re
import django

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'mdh_intranet.settings')
django.setup()

from mdh_intranet.sop_manual.models import SOP, SOPCategory
from django.contrib.auth.models import User

file_path = r"c:\Users\HomePC\mdh_intranet\media\documents\2026\02\MDH_SOP_Manual_Full_Content.docx"

# Mapping of SOP code prefixes to categories
CATEGORY_MAPPING = {
    'CLIN': 'Clinical Procedures',
    'NUR': 'Clinical Procedures',
    'SAF': 'Safety & Infection Control',
    'INF': 'Safety & Infection Control',
    'PAT': 'Patient Care',
    'EME': 'Emergency Protocols',
    'EMG': 'Emergency Protocols',
    'ADM': 'Administrative',
    'FIN': 'Administrative',
    'OPS': 'Administrative',
    'HR': 'Administrative',
    'QUA': 'Quality Assurance',
    'QA': 'Quality Assurance',
}

def categorize_sop(code):
    if code:
        match = re.match(r'MDH-([A-Z]+)-', code)
        if match:
            prefix = match.group(1)
            return CATEGORY_MAPPING.get(prefix, 'Administrative')
    return 'Administrative'

def extract_and_import_sops():
    print("Loading document...")
    doc = Document(file_path)
    print(f"Document loaded. Total paragraphs: {len(doc.paragraphs)}")
    
    # Extract all text first (faster)
    print("Extracting all text...")
    all_paras = [p.text for p in doc.paragraphs]
    print(f"Text extracted. Processing...")
    
    # Find all SOP headers
    # Support various dash types: -, –, —
    sop_headers = []
    for i, text in enumerate(all_paras):
        # Broad lookup for MDH-...
        if 'MDH-' in text:
            # Match code and title
            match = re.search(r'(MDH-[A-Z]{2,4}-\d{3})\s*[-–—]\s*(.+)', text)
            if match:
                sop_headers.append({
                    'index': i,
                    'code': match.group(1),
                    'title': match.group(2).strip()
                })
    
    print(f"Found {len(sop_headers)} SOP headers\n")
    
    if not sop_headers:
        print("No headers found with regex. Trying fallback matching...")
        # Fallback: any line starting with MDH-
        for i, text in enumerate(all_paras):
            if text.startswith('MDH-'):
                parts = text.split(' ', 1)
                code = parts[0].strip(' —-–')
                title = parts[1].strip(' —-–') if len(parts) > 1 else "Unknown Title"
                sop_headers.append({
                    'index': i,
                    'code': code,
                    'title': title
                })
        print(f"Found {len(sop_headers)} headers with fallback\n")

    # For each SOP, gather content
    sops_data = []
    for i, header in enumerate(sop_headers):
        start_idx = header['index'] + 1
        end_idx = sop_headers[i+1]['index'] if i+1 < len(sop_headers) else len(all_paras)
        
        # Collect content
        content_paras = [all_paras[j].strip() for j in range(start_idx, end_idx) if all_paras[j].strip()]
        
        # Avoid collecting the next header if regex found it inside a paragraph
        # But since we use indices, this should be fine.
        
        content = '\n\n'.join(content_paras)
        
        # Relax content length requirement for testing
        if len(content) > 10:
            sops_data.append({
                'code': header['code'],
                'title': header['title'],
                'content': content
            })
        else:
            print(f"DEBUG: Skipping {header['code']} - content length {len(content)}")
        
        if (i+1) % 50 == 0:
            print(f"Processed {i+1}/{len(sop_headers)} headers...")
    
    print(f"\nExtracted {len(sops_data)} SOPs with content\n")
    
    # Import
    user = User.objects.first()
    if not user:
        print("No users found.")
        return
    
    imported_count = 0
    existing_count = 0
    
    print("Beginning database import...")
    for sop_data in sops_data:
        cat_name = categorize_sop(sop_data['code'])
        category = SOPCategory.objects.filter(name=cat_name).first()
        
        if not category:
            # Fallback to Administrative if not found
            category = SOPCategory.objects.filter(name='Administrative').first()
            if not category:
                continue
        
        full_title = f"{sop_data['code']} - {sop_data['title']}"
        if len(full_title) > 200:
            full_title = full_title[:197] + "..."
        
        content = sop_data['content']
        if len(content) > 50000:
            content = content[:50000] + "\n\n[Content truncated...]"
        
        sop, created = SOP.objects.get_or_create(
            title=full_title,
            defaults={
                'category': category,
                'content': content,
                'version': '1.0',
                'status': 'Published',
                'created_by': user
            }
        )
        
        if created:
            imported_count += 1
            if imported_count <= 10 or imported_count % 20 == 0:
                print(f"✓ {imported_count}. {sop.title[:70]}")
        else:
            existing_count += 1
    
    print(f"\n{'='*70}")
    print(f"✅ Import Finish!")
    print(f"{'='*70}")
    print(f"Successfully imported: {imported_count} SOPs")
    print(f"Already existed: {existing_count} SOPs")
    print(f"{'='*70}")

if __name__ == '__main__':
    extract_and_import_sops()
