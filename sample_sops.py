
import os
import django
from docx import Document as DocxDocument
import re

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'mdh_intranet.settings')
django.setup()
from mdh_intranet.documents.models import Document

def sample_sops(doc_id):
    d = Document.objects.get(id=doc_id)
    doc = DocxDocument(d.file.path)
    
    # Let's check the first one (usually cleanest)
    print("--- SOP 1 Sample (FIN-001) ---")
    for i in range(180, 250):
        print(f"[{i}] {doc.paragraphs[i].text.strip()}")
    
    # Let's check a middle one (might be messier)
    print("\n--- SOP 50 Sample ---")
    # Finding header for CLIN-001 or similar
    found = 0
    for i, p in enumerate(doc.paragraphs):
        if re.search(r'MDH-CLIN-\d+', p.text):
            found += 1
            if found == 5: # Fifth clinical SOP
                print(f"Found at {i}: {p.text}")
                for j in range(i, i+50):
                    if j < len(doc.paragraphs):
                        print(f"[{j}] {doc.paragraphs[j].text.strip()}")
                break

if __name__ == "__main__":
    sample_sops(4)
