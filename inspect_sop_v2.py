
from docx import Document
import os
import re

file_path = r"c:\Users\HomePC\mdh_intranet\media\documents\2026\02\MDH_SOP_Manual_Full_Content.docx"

doc = Document(file_path)

# Let's try a different approach - look at headings more carefully
print("Analyzing document structure...")
print(f"Total paragraphs: {len(doc.paragraphs)}\n")

# Sample different parts of the document
indices = [0, 100, 500, 1000, 2000, 5000, 10000, 15000]
for idx in indices:
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        style_name = p.style.name if p.style else "No style"
        print(f"Para {idx}: [{style_name}] {p.text[:150]}")
        
print("\n\nLooking for heading paragraphs...")
heading_count = 0
for i, p in enumerate(doc.paragraphs[:5000]):  # First 5000 paragraphs
    if p.style and p.style.name and 'Heading' in p.style.name:
        heading_count += 1
        if heading_count <= 30:  # Show first 30 headings
            print(f"{i}: [{p.style.name}] {p.text[:100]}")
