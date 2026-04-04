
from docx import Document
import os
import re
import django

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'mdh_intranet.settings')
django.setup()

from mdh_intranet.sop_manual.models import SOP, SOPCategory
from django.contrib.auth.models import User

file_path = r"c:\Users\HomePC\mdh_intranet\media\documents\2026\02\MDH_SOP_Manual_Full_Content.docx"

# Mapping of SOP code prefixes to categories
CATEGORY_MAPPING = {
    'CLIN': 'Clinical Procedures',
    'NUR': 'Clinical Procedures',
    'SAF': 'Safety & Infection Control',
    'INF': 'Safety & Infection Control',
    'PAT': 'Patient Care',
    'EME': 'Emergency Protocols',
    'EMG': 'Emergency Protocols',
    'ADM': 'Administrative',
    'FIN': 'Administrative',
    'OPS': 'Administrative',
    'HR': 'Administrative',
    'QUA': 'Quality Assurance',
    'QA': 'Quality Assurance',
}

def categorize_sop(code):
    if code:
        match = re.match(r'MDH-([A-Z]+)-', code)
        if match:
            prefix = match.group(1)
            return CATEGORY_MAPPING.get(prefix, 'Administrative')
    return 'Administrative'

def extract_and_import_sops():
    doc = Document(file_path)
    
    # Find all paragraphs with MDH codes (these are the headers)
    sop_headers = []
    for i, para in enumerate(doc.paragraphs):
        match = re.search(r'(MDH-[A-Z]{2,4}-\d{3})\s*[—-]\s*(.+)', para.text)
        if match:
            sop_headers.append({
                'index': i,
                'code': match.group(1),
                'title': match.group(2).strip(),
                'full_text': para.text.strip()
            })
    
    print(f"Found {len(sop_headers)} SOP headers\n")
    
    # For each SOP, gather content from its paragraph until the next SOP header
    sops_data = []
    for i, header in enumerate(sop_headers):
        start_idx = header['index'] + 1  # Start after the header
        end_idx = sop_headers[i+1]['index'] if i+1 < len(sop_headers) else len(doc.paragraphs)
        
        # Collect content paragraphs
        content_paras = []
        for j in range(start_idx, end_idx):
            text = doc.paragraphs[j].text.strip()
            if text:  # Only add non-empty paragraphs
                content_paras.append(text)
        
        content = '\n\n'.join(content_paras)
        
        if len(content) > 50:  # Only if has meaningful content
            sops_data.append({
                'code': header['code'],
                'title': header['title'],
                'content': content
            })
    
    print(f"Extracted {len(sops_data)} SOPs with content\n")
    
    # Show sample
    for i, sop in enumerate(sops_data[:10]):
        print(f"{i+1}. {sop['code']}: {sop['title'][:60]}")
        print(f"   Content: {len(sop['content'])} chars\n")
    
    # Import
    user = User.objects.first()
    if not user:
        print("No users found.")
        return
    
    imported_count = 0
    
    for sop_data in sops_data:
        cat_name = categorize_sop(sop_data['code'])
        category = SOPCategory.objects.filter(name=cat_name).first()
        
        if not category:
            print(f"⚠ Skipping {sop_data['code']}: Category '{cat_name}' not found")
            continue
        
        full_title = f"{sop_data['code']} - {sop_data['title']}"
        if len(full_title) > 200:
            full_title = full_title[:197] + "..."
        
        content = sop_data['content']
        if len(content) > 50000:
            content = content[:50000] + "\n\n[Content truncated...]"
        
        sop, created = SOP.objects.get_or_create(
            title=full_title,
            defaults={
                'category': category,
                'content': content,
                'version': '1.0',
                'status': 'Published',
                'created_by': user
            }
        )
        
        if created:
            imported_count += 1
            if imported_count <= 10 or imported_count % 20 == 0:
                print(f"✓ {imported_count}. {sop.title[:80]}")
    
    print(f"\n\n{'='*70}")
    print(f"✅ Import Complete!")
    print(f"{'='*70}")
    print(f"Successfully imported: {imported_count} SOPs")
    print(f"{'='*70}")

if __name__ == '__main__':
    extract_and_import_sops()
