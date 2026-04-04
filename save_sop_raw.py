
import os
import django
from docx import Document as DocxDocument
import re
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'mdh_intranet.settings')
django.setup()
from mdh_intranet.documents.models import Document

def extract_full_sop_to_temp(doc_id, start_idx):
    d = Document.objects.get(id=doc_id)
    doc = DocxDocument(d.file.path)
    content = []
    
    elements = list(doc.element.body.iterchildren())
    
    with open('sop_raw_text.txt', 'w', encoding='utf-8') as f:
        for i in range(start_idx, len(elements)):
            element = elements[i]
            if isinstance(element, CT_P):
                p = Paragraph(element, doc)
                text = p.text.strip()
                if i > start_idx + 10 and re.search(r'MDH-[A-Z]+-\d+', text) and len(text) < 100:
                    break
                if text:
                    f.write(f"P: {text}\n")
            elif isinstance(element, CT_Tbl):
                t = Table(element, doc)
                data = [[c.text.strip() for c in r.cells] for r in t.rows]
                f.write(f"T: {data}\n")

if __name__ == "__main__":
    d = Document.objects.get(id=4)
    doc = DocxDocument(d.file.path)
    target_p = doc.paragraphs[182]
    elements = list(doc.element.body.iterchildren())
    start_elm_idx = -1
    for i, el in enumerate(elements):
        if isinstance(el, CT_P):
            if Paragraph(el, doc).text == target_p.text:
                start_elm_idx = i
                break
    
    if start_elm_idx != -1:
        extract_full_sop_to_temp(4, start_elm_idx)
